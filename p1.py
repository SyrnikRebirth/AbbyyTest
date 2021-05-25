import random
import docx
import subprocess

sample = " sample text "
document = docx.Document()

document.add_paragraph(sample * random.randrange(10))

document.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

document.add_paragraph(sample * random.randrange(10))

document.paragraphs[1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

document.add_picture('cat.jpeg', width = docx.shared.Cm(15))

document.paragraphs[2].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

table = document.add_table(rows = 2, cols = 3)
for i in range(2):
    for j in range(3):
        cell = table.cell(i, j)
        cell.text = sample

document.save('test.docx')

subprocess.run(['libreoffice', '--convert-to', 'pdf' ,'test.docx'])

subprocess.run(['rm', 'test.docx'])
